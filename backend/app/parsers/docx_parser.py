from docx import Document as DocxDocument

from app.parsers.base import BaseParser, ParseResult


class DOCXParser(BaseParser):
    def parse(self, file_path: str) -> ParseResult:
        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return ParseResult(raw_text=text, metadata={"paragraph_count": len(doc.paragraphs)})
